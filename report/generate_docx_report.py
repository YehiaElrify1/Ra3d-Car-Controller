import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_rtl_paragraph(doc, text="", style=None):
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    if text:
        add_rtl_run(p, text)
    return p

def add_rtl_run(paragraph, text, bold=False, italic=False, size_pt=11, font_name='Arial', color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    return run

def add_rtl_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    set_rtl(h)
    size_pt = 18 if level == 1 else (14 if level == 2 else 12)
    color = RGBColor(0, 51, 102) if level == 1 else RGBColor(51, 102, 153)
    add_rtl_run(h, text, bold=True, size_pt=size_pt, color=color)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    
    # Set light gray shading
    shading = OxmlElement('w:shd')
    shd_val = qn('w:fill')
    shading.set(shd_val, 'F4F4F4')
    cell._tc.get_or_add_tcPr().append(shading)
    
    # Remove borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = paragraph.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is not None:
        pPr.remove(bidi)

def add_centered_image(doc, img_path, width_inches=4.5, caption_text=""):
    if not os.path.exists(img_path):
        print(f"Warning: image not found at {img_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    
    if caption_text:
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(12)
        add_rtl_run(caption_p, caption_text, italic=True, size_pt=9.5, color=RGBColor(100, 100, 100))

def read_file(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return f"Error reading file: {e}"

def main():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(12)
    add_rtl_run(title_p, "التقرير الفني الشامل لمشروع التخرج", bold=True, size_pt=26, color=RGBColor(0, 51, 102))

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    add_rtl_run(subtitle_p, "نظام التحكم والمراقبة للسيارة الذكية رعد (Ra3d Car Controller)\nتطبيق هاتف محمول متعدد المنصات (.NET MAUI) ونظام مدمج (ATmega32) عبر البلوتوث", size_pt=14, color=RGBColor(51, 102, 153))

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(120)
    add_rtl_run(info_p, "مادة: البرمجة المتقدمة (Advanced Programming)\n", bold=True, size_pt=12)
    add_rtl_run(info_p, "قسم هندسة الحاسبات والتحكم - السنة الثالثة\n", size_pt=12)
    add_rtl_run(info_p, "العام الدراسي: 2025/2026", size_pt=12)
    
    doc.add_page_break()

    # 2. Introduction
    add_rtl_heading(doc, "1. مقدمة عامة عن المشروع (Introduction)", level=1)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "يمثل هذا المشروع نظاماً متكاملاً يربط بين تقنيات الأنظمة المدمجة (Embedded Systems) وتطوير تطبيقات الموبايل المتقدمة باستخدام إطار العمل الحديث .NET MAUI. يهدف المشروع إلى تصميم وبناء سيارة ذكية لاسلكية يُطلق عليها اسم ")
    add_rtl_run(p, "رعد (Ra3d Car)", bold=True)
    add_rtl_run(p, "، حيث يتم التحكم في حركتها بشكل كامل ومراقبة مؤشراتها الحيوية لاسلكياً عبر البلوتوث الكلاسيكي (Bluetooth Classic).")

    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "يعتمد العتاد الصلب (Hardware) للسيارة على لوحة ميكروكنترولر من نوع ")
    add_rtl_run(p, "ATmega32", bold=True)
    add_rtl_run(p, "، والتي تمثل العقل المدبر المسؤول عن معالجة الإشارات، وتوليد إشارات تعديل عرض النبضة (PWM) لتشغيل المحركات، وقراءة الحساسات. وتتواصل مع تطبيق الموبايل عبر وحدة البلوتوث ")
    add_rtl_run(p, "HC-05", bold=True)
    add_rtl_run(p, " باستخدام بروتوكول المنافذ التسلسلية (SPP - Serial Port Profile). بينما تم بناء تطبيق الموبايل ليعمل بكفاءة عالية على منصات متعددة، ويوفر واجهة قيادة ومراقبة حيوية بطابع نيون عصري، ويستعرض بيانات الحساسات مثل درجة الحرارة والمسافة بشكل فوري.")

    # 3. System Architecture
    add_rtl_heading(doc, "2. بنية النظام العامة (System Architecture)", level=1)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "تتكون بنية النظام الإجمالية من اتجاهين رئيسيين للبيانات:")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "1. اتجاه التحكم (Control Flow): ", bold=True)
    add_rtl_run(p, "يقوم المستخدم بالضغط على أزرار التوجيه في واجهة القيادة بالتطبيق، فيقوم التطبيق بإرسال أحرف ASCII مخصصة ('F', 'B', 'R', 'L') تمثل الاتجاهات، وحرف 'S' للتوقف الفوري عند رفع الإصبع. يستقبل موديول HC-05 هذه الإشارات ويمررها للمعالج عبر واجهة UART، ليقوم الأخير بضبط قيم منافذ التحكم المتصلة بمشغل المحركات L293D لتوجيه السيارة.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "2. اتجاه المراقبة (Telemetry Flow): ", bold=True)
    add_rtl_run(p, "يقوم الميكروكنترولر في حلقة مستمرة بقراءة الحساسات المتصلة به (حساس الأمواج فوق الصوتية HC-SR04 وحساس درجة الحرارة LM35) ومن ثم تحويل القراءات إلى قيم فعلية. يتم إرسال هذه البيانات في حزمة نصية منسقة بصيغة 'D:xx,T:yy#' عبر البلوتوث إلى التطبيق. يقوم التطبيق بقراءة هذه البيانات باستمرار، وتفكيك الحزم البرمجية وتحديث عناصر واجهة المستخدم رسومياً.")

    # 4. Software Stack and Libraries
    add_rtl_heading(doc, "3. المكونات البرمجية والمكتبات المستخدمة (Libraries & APIs)", level=1)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "تم بناء تطبيق الموبايل بالاعتماد على أحدث التقنيات البرمجية لضمان سرعة الاستجابة وجمالية التصميم، وتتضمن المكونات:")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "* إطار العمل .NET MAUI: ", bold=True)
    add_rtl_run(p, "أحدث إطار عمل من مايكروسوفت لبناء تطبيقات أصلية لمنصات متعددة باستخدام كود واحد بلغة C# وواجهات XAML. يتيح لنا هذا النظام تصميم واجهات مستخدم متجاوبة تماماً وتصدير التطبيق لنظام أندرويد.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "* مكتبة InTheHand.Net.Bluetooth (32feet.net): ", bold=True)
    add_rtl_run(p, "مكتبة مفتوحة المصدر تم دمجها عبر NuGet Package لإدارة اتصالات البلوتوث الكلاسيكي في بيئة دوت نت. تكمن أهميتها في توفير كلاسات موحدة مثل BluetoothClient للاتصال وقناة البث التسلسلي.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "* واجهة برمجة أندرويد الأصلية (Android.Bluetooth API): ", bold=True)
    add_rtl_run(p, "نظراً لأن أنظمة أندرويد الحديثة تفرض قيوداً صارمة على البحث والاتصال بالبلوتوث وتتطلب استخدام صلاحيات خاصة بـ Runtime وتخصيص قنوات اتصال RFCOMM فريدة، تم تضمين كود برمجي مخصص لمنصة أندرويد باستخدام التوجيه الشرطي #if ANDROID. يتيح هذا الكود التخاطب المباشر مع نظام تشغيل الهاتف لفتح قنوات اتصال مستقرة خالية من الكراش.")

    # 5. App.xaml & App.xaml.cs
    add_rtl_heading(doc, "4. ملف التهيئة وإدارة الحالة العامة: App.xaml و App.xaml.cs", level=1)
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "ملف ")
    add_rtl_run(p, "App.xaml", bold=True)
    add_rtl_run(p, " هو الملف التعريفي الرئيسي للتطبيق حيث يقوم بدمج ملفات الألوان والتنسيقات الجمالية للواجهة. فيما يلي كود الملف كاملاً:")
    add_code_block(doc, read_file("../App.xaml"))
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "بينما يمثل ملف ")
    add_rtl_run(p, "App.xaml.cs", bold=True)
    add_rtl_run(p, " العقل المدبر للتطبيق حيث يحتوي على المتغيرات والأحداث والدوال العامة المتاحة لكل شاشات التطبيق لضمان الاتصال المشترك وتحديث الحالة فورا:")
    add_code_block(doc, read_file("../App.xaml.cs"))

    add_rtl_heading(doc, "تحليل الكود جزء بجزء وتأثيره على التطبيق:", level=2)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "1. الحدث `ConnectionChanged` والمتغير `IsConnected`: ", bold=True)
    add_rtl_run(p, "يمثل هذا الجزء تطبيقاً لنمط التصميم القائم على الأحداث (Event-Driven Architecture). عند نجاح الاتصال أو انقطاعه، يتم تعديل قيمة المتغير `IsConnected`. يقوم الـ Setter الخاص به بإطلاق حدث `ConnectionChanged?.Invoke()` تلقائياً، مما ينبه جميع الصفحات المشتركة بالحدث فوراً لتحديث الألوان والنصوص (مثل تحول مؤشر حالة الاتصال للأخضر أو الأحمر).")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "2. قناة البث `BthStream`: ", bold=True)
    add_rtl_run(p, "تيار عام مشترك من نوع Stream يمثل القناة الفيزيائية لنقل البيانات. تتم القراءة والكتابة من خلاله من مختلف شاشات التطبيق دون الحاجة لإعادة إنشاء قنوات اتصال جديدة.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "3. الحدث `TemperatureChanged` ودالة `UpdateTemperature`: ", bold=True)
    add_rtl_run(p, "تستقبل هذه الدالة درجة الحرارة اللحظية المستلمة من الحساس في الخلفية، وتقوم بتخزينها في `CurrentTemperature` وإطلاق حدث `TemperatureChanged` لنقل القيمة مباشرة إلى شاشة التحليلات ديناميكياً.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "4. دالة إرسال الأوامر `SendCommand`: ", bold=True)
    add_rtl_run(p, "تتحقق أولاً من صلاحية الاتصال وقابلية الكتابة على القناة التسلسلية، ثم تقوم بتحويل النص إلى بايتات ASCII وإرسالها بشكل غير متزامن `BthStream.WriteAsync` لضمان عدم توقف الواجهة. وفي حال فشل الإرسال (انقطاع الاتصال)، تقوم فوراً بضبط `IsConnected = false` لحماية التطبيق من الانهيار.")

    # 6. SetupPage.xaml & SetupPage.xaml.cs
    doc.add_page_break()
    add_rtl_heading(doc, "5. واجهة البحث وإعداد الاتصال: SetupPage.xaml و SetupPage.xaml.cs", level=1)
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "ملف ")
    add_rtl_run(p, "SetupPage.xaml", bold=True)
    add_rtl_run(p, " يحدد واجهة المستخدم الرسومية لعملية البحث. يتميز بتصميمه الداكن بطابع نيون ويحتوي على CollectionView لعرض قائمة الأجهزة المكتشفة، ومؤشر نشاط ActivityIndicator لإظهار جاري البحث، وأزرار لبدء البحث أو قطع الاتصال:")
    add_code_block(doc, read_file("../SetupPage.xaml"))
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "بينما يحتوي ملف الكود الخلفي ")
    add_rtl_run(p, "SetupPage.xaml.cs", bold=True)
    add_rtl_run(p, " على منطق فحص وصلاحيات التشغيل للبلوتوث والموقع الجغرافي، وإجراء عملية المسح عبر البث المباشر لأندرويد، وتأمين الاتصال التسلسلي:")
    add_code_block(doc, read_file("../SetupPage.xaml.cs"))

    add_rtl_heading(doc, "تحليل الكود جزء بجزء وتأثيره على التطبيق:", level=2)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "1. كلاس الصلاحيات المخصص `BluetoothPermissions`: ", bold=True)
    add_rtl_run(p, "في أنظمة أندرويد الحديثة (بدءاً من Android 12)، يتطلب النظام صلاحيات تشغيل متقدمة هي BLUETOOTH_SCAN و BLUETOOTH_CONNECT. يقوم هذا الكلاس بالوراثة من BasePlatformPermission وتوفير الصلاحيات المطلوبة ديناميكياً ليطلبها التطبيق في وقت التشغيل لمنع الكراش الأمني.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "2. دالة بدء البحث `OnScanClicked`: ", bold=True)
    add_rtl_run(p, "تقوم بطلب الصلاحيات أولاً. وعند الموافقة، تقوم بفحص الأجهزة المقترنة مسبقاً (Bonded Devices) وإضافتها فوراً لتوفير الوقت على المستخدم. ثم تقوم بإنشاء مستقبل البث المباشر `BluetoothDiscoveryReceiver` وتسجيله في نظام أندرويد ليلتقط إشعارات الأجهزة القريبة الفعالة لحظياً عبر `StartDiscovery()`.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "3. دالة الاختيار وتأسيس الاتصال `OnDeviceSelected`: ", bold=True)
    add_rtl_run(p, "عند نقر المستخدم على جهاز من القائمة (مثل HC-05)، تسأله لتأكيد الرغبة بالاتصال. ثم تبدأ عملية الاتصال بشكل غير متزامن. في بيئة أندرويد، نستخدم مقبس RFCOMM الأصلي المرتبط بالمعرف الموحد الافتراضي للمنافذ التسلسلية SPP UUID. يضمن تشغيل الدالة داخل `Task.Run` استمرارية سلاسة الواجهة أثناء محاولة المصافحة (Handshake). وعند النجاح، يتم تغليف قنوات الاتصال بـ `BluetoothSocketStream` وضبط التطبيق كمتصل.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "4. الكلاس المساعد `BluetoothSocketStream`: ", bold=True)
    add_rtl_run(p, "كلاس يرث من Stream ويقوم بحل مشكلة تقاطع قنوات الإدخال والإخراج في نظام أندرويد عند استخدام المقابس الأصلية، حيث يتيح القراءة والكتابة المتزامنة بأمان دون حدوث تضارب في البيانات.")

    # Image setup page
    add_rtl_heading(doc, "لقطة شاشة توضيحية لواجهة إعداد الاتصال والبحث:", level=2)
    add_centered_image(doc, "setup_screenshot.png", width_inches=4.2, caption_text="شكل (1): واجهة رادار البلوتوث المخصصة للبحث والاقتران بموديل HC-05")

    # 7. DrivePage.xaml & DrivePage.xaml.cs
    doc.add_page_break()
    add_rtl_heading(doc, "6. واجهة القيادة ورادار الأمان اللحظي: DrivePage.xaml و DrivePage.xaml.cs", level=1)
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "ملف ")
    add_rtl_run(p, "DrivePage.xaml", bold=True)
    add_rtl_run(p, " يحدد واجهة التحكم بالسيارة. يحتوي على مؤشر حالة الاتصال العلوي، ومؤشر رادار دائري يتغير لونه ديناميكياً لعرض المسافة الفاصلة عن العوائق، ولوحة تحكم حركية مكونة من أزرار اتجاهية تعمل باللمس المستمر:")
    add_code_block(doc, read_file("../DrivePage.xaml"))
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "ملف الكود الخلفي ")
    add_rtl_run(p, "DrivePage.xaml.cs", bold=True)
    add_rtl_run(p, " يربط أحداث لمس الأزرار بإرسال الأوامر للسيارة، ويحتوي على حلقة الاستماع بالخلفية وتفكيك بيانات الحساسات القادمة من الميكروكنترولر وتحديث مؤشرات الرادار والألوان:")
    add_code_block(doc, read_file("../DrivePage.xaml.cs"))

    add_rtl_heading(doc, "تحليل الكود جزء بجزء وتأثيره على التطبيق:", level=2)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "1. الاشتراك بحدث الاتصال `App.ConnectionChanged`: ", bold=True)
    add_rtl_run(p, "عند ظهور الشاشة، تشترك في الحدث العام. وعند إطلاقه، تقوم دالة `UpdateConnectionStatus` بتحديث مؤشر الحالة الرسومي باللون الأخضر بكلمة 'Connected' وبدء الاستماع الفوري للبيانات، أو باللون الأحمر وبكلمة 'Disconnected' وإعادة تهيئة مؤشرات الرادار لوضع الاستعداد.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "2. أحداث أزرار التوجيه `Pressed` و `Released`: ", bold=True)
    add_rtl_run(p, "تمثل جوهر التحكم الحركي. عند ضغط زر للأمام، يرسل فوراً وبشكل مستمر الحرف 'F'، وعند رفعه يرسل الحرف 'S'. تضمن هذه الآلية أقصى درجات الأمان حركياً للسيارة، فلو انقطع الاتصال أو أزال المستخدم إصبعه، تتوقف السيارة فوراً.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "3. حلقة الاستماع بالخلفية `StartListeningToCarData`: ", bold=True)
    add_rtl_run(p, "تعمل بشكل منفصل وغير متزامن. تستمر في قراءة تيار البلوتوث عبر `ReadAsync` وتجميع القطع النصية المستلمة في متغير مؤقت. وتقوم بحلقة فحص للبحث عن رمز النهاية '#'؛ وفور إيجاده، تقتطع الرسالة الكاملة وتقسمها باستخدام الفاصلة ',' لفصل قراءة المسافة عن درجة الحرارة.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "4. معالجة وتحديث واجهة الرادار `UpdateRadarUI`: ", bold=True)
    add_rtl_run(p, "تقوم بتفسير المسافة المقاسة بالسنتيمتر: إذا كانت المسافة آمنة تتلون الواجهة بالأخضر وتظهر 'Radar: Clear Path'. وإذا اقتربت السيارة من عائق (بين 15 و 30 سم) يتحول الرادار للأصفر كلون تحذيري. أما إذا اقتربت بشدة (أقل من 15 سم) تضيء الشاشة بالأحمر التحذيري وتظهر عبارة 'OBSTACLE DETECTED!'.")

    # Image drive page
    add_rtl_heading(doc, "لقطة شاشة توضيحية لواجهة التحكم والقيادة الفعالة:", level=2)
    add_centered_image(doc, "drive_screenshot.png", width_inches=4.2, caption_text="شكل (2): واجهة القيادة وتظهر حالة الاتصال النشط ومؤشر الرادار الآمن ولوحة التوجيه")

    # 8. AnalyticsPage.xaml & AnalyticsPage.xaml.cs
    doc.add_page_break()
    add_rtl_heading(doc, "7. واجهة المراقبة والتحليلات الحرارية: AnalyticsPage.xaml و AnalyticsPage.xaml.cs", level=1)
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "ملف ")
    add_rtl_run(p, "AnalyticsPage.xaml", bold=True)
    add_rtl_run(p, " يحدد واجهة عرض درجة حرارة محركات السيارة وبطاريتها. يعرض درجة الحرارة بخط كبير بارز داخل مؤشر حراري دائري، مع وجود بطاقة لعرض إشعارات وتدابير الأمان المقترحة تلقائياً:")
    add_code_block(doc, read_file("../AnalyticsPage.xaml"))
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "ملف الكود الخلفي ")
    add_rtl_run(p, "AnalyticsPage.xaml.cs", bold=True)
    add_rtl_run(p, " يشترك في حدث تغير الحرارة العام، ويقوم بتطبيق منطق الحماية الحرارية عبر تصنيف مستويات الخطورة وتغيير ألوان النصوص والإنذارات ديناميكياً:")
    add_code_block(doc, read_file("../AnalyticsPage.xaml.cs"))

    add_rtl_heading(doc, "تحليل الكود جزء بجزء وتأثيره على التطبيق:", level=2)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "1. دالة التحديث الفوري `OnAppearing`: ", bold=True)
    add_rtl_run(p, "عندما ينتقل المستخدم إلى تبويب التحليلات، تضمن هذه الدالة استرداد آخر درجة حرارة مسجلة ومحفوظة وعرضها فوراً بدلاً من انتظار وصول قراءة جديدة من البلوتوث.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "2. آلية الاستقبال الآمن `MainThread.BeginInvokeOnMainThread`: ", bold=True)
    add_rtl_run(p, "نظراً لأن عملية الاستماع للبلوتوث تتم في خيط معالجة خلفي (Background Thread)، فإن محاولة تحديث واجهة المستخدم بشكل مباشر من هذا الخيط ستؤدي لنهيار التطبيق (Crash). تقوم هذه التعليمة بنقل عملية التحديث بأمان إلى خيط المعالجة الرئيسي للواجهات.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "3. منطق الحماية الحرارية الثلاثي `UpdateThermalUI`: ", bold=True)
    add_rtl_run(p, "يقوم بفحص الحرارة الواردة وتحديد مستوى الخطر: إذا كانت الحرارة طبيعية (أقل من 40) يظهر باللون الأخضر للإعلان عن سلامة المحركات وتوازن الحمل. وإذا ارتفعت (بين 40 و 55) يضيء بالأصفر لتنبيه المستخدم بارتفاع الحمل طفيفاً ومراقبة الوضع. أما إذا تجاوزت 55 درجة، فيعتبرها التطبيق حالة طوارئ حرجة لتجنب احتراق الملفات أو انفجار الخلايا الكيميائية للبطارية، ويطلق تحذيراً شديداً باللون الأحمر ويطلب إيقاف السيارة فوراً للتبريد.")

    # Image analytics page
    add_rtl_heading(doc, "لقطة شاشة توضيحية لواجهة التحليلات الحرارية والإنذارات:", level=2)
    add_centered_image(doc, "analytics_screenshot.png", width_inches=4.2, caption_text="شكل (3): واجهة المراقبة وتظهر قراءات الحساسات الحرارية ومستوى الأمان واستقرار الحمل")

    # 9. carmain.c (ATmega32 C Firmware)
    doc.add_page_break()
    add_rtl_heading(doc, "8. كود نظام التحكم والقيادة الذاتي للميكروكنترولر: carmain.c", level=1)
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "يمثل كود لغة C المكتوب لميكروكنترولر ")
    add_rtl_run(p, "ATmega32", bold=True)
    add_rtl_run(p, " المحرك الأساسي للسيارة. يقوم الكود بإدارة الاتصال التسلسلي مع البلوتوث، وقراءة إشارة حساس الموجات فوق الصوتية والجهد التماثلي لحساس الحرارة، وإدارة محركات الحركة مع نظام أمان ذاتي لتفادي التصادم في حال اكتشاف عائق قريب دون الحاجة لتدخل تطبيق الموبايل. الكود بالكامل موضح أدناه:")
    add_code_block(doc, read_file("../final_car_project_v1.X/carmain.c"))

    add_rtl_heading(doc, "تحليل الكود جزء بجزء وتأثيره على السيارة والتطبيق:", level=2)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "1. محرك الاتصال التسلسلي `UART_Init` و `UART_Receive`: ", bold=True)
    add_rtl_run(p, "يقوم بضبط معدل الباود ريت على 9600 وتفعيل الإرسال والاستقبال متزامناً مع موديول HC-05. الدالة `UART_Receive` تفحص السجل UCSRA للتحقق من وصول حرف جديد؛ وفي حال وصوله تقوم بإرجاع الحرف المستلم فوراُ دون حظر المعالج.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "2. تهيئة وقراءة حساس الحرارة `ADC_Read`: ", bold=True)
    add_rtl_run(p, "يتم قراءة الحساس LM35 الموصل بالقناة السادسة. يُنتج الحساس 10 ملي فولت لكل درجة حرارة واحدة. يقوم المتحكم بتحويل القراءة التماثلية لرقمية وحساب الدرجة المئوية المقابلة عبر المعادلة الرياضية `temp_c = ADC * 500 / 1024` والتي تعادل قيمة الجهد الفعلي بالملي فولت مقسوماً على 10.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "3. قياس المسافة وتفادي التعليق `Get_Distance`: ", bold=True)
    add_rtl_run(p, "يقوم بتوليد نبضة قدح Trig مدتها 10 ميكروثانية لحساس الأمواج فوق الصوتية، ثم قياس عرض نبضة الصدى Echo. تم تضمين مؤقت أمان (Timeout) داخل كود القياس لحماية البرنامج من التعليق اللانهائي في حال عدم عودة الصدى (مثل توجيه الحساس للسماء المفتوحة)، حيث يعيد القيمة 999 للدلالة على أمان الطريق تماماً.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "4. نظام تفادي الاصطدام التلقائي (Emergency Collision Avoidance): ", bold=True)
    add_rtl_run(p, "يمثل هذا الكود حلقة الأمان في السيارة. يتم فحص المسافة المقاسة في كل دورة: إذا كانت المسافة أقل من 15 سم، يتدخل الميكروكنترولر فوراً ويقوم بتصفير مخارج التحكم بالمحركات `PORTA &= 0xF0` لإيقافها إجبارياً وتجاهل أوامر المستخدم لتجنب الحوادث، مما يعطي السيارة ذكاءً ذاتياً لحماية هيكلها.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "5. فك شفرة الأوامر وتوجيه الحركة: ", bold=True)
    add_rtl_run(p, "تتحقق حلقة الاختيار Switch-Case من الحرف الوارد: الحرف 'F' يوجه القطبية لتشغيل المحركين للأمام، الحرف 'B' للخلف، بينما يقوم الحرفان 'R' و 'L' بتشغيل أحد المحركات للأمام والآخر للخلف لإحداث دوران مغزلي سريع حول مركز السيارة (Zero-Radius Spin) مما يعطي السيارة مرونة فائقة في الأماكن الضيقة. والحرف 'S' يقطع الطاقة عن المحركات.")

    # 10. Hardware Connections & Schematics
    doc.add_page_break()
    add_rtl_heading(doc, "9. المخططات الهندسية وتوصيل الدائرة (Hardware Connections)", level=1)
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "يشتمل النظام على لوحة توصيلات هندسية متكاملة لربط ميكروكنترولر ATmega32 مع الحساسات ومنافذ الاتصال والتحكم بالمحركات. التوصيلات موضحة كالتالي:")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "* موديول البلوتوث HC-05: ", bold=True)
    add_rtl_run(p, "موصل بالبورت التسلسلي للمتحكم (RX/PD0 مع TX للموديول) و (TX/PD1 مع RX للموديول عبر مقسم جهد 1K/2K لحماية مدخل الموديول).")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "* حساس المسافة HC-SR04: ", bold=True)
    add_rtl_run(p, "موصل بالطرفين PA4 (Trig) كخرج والطرف PA5 (Echo) كدخل رقمي مقاس بالزمن.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "* حساس الحرارة LM35: ", bold=True)
    add_rtl_run(p, "موصل بالمدخل التماثلي PA6 كقناة دخل للمحول ADC.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "* مشغل المحركات L293D ومحركات الحركة: ", bold=True)
    add_rtl_run(p, "موصل بأطراف الاتجاه PA0, PA1, PA2, PA3 وأطراف التحكم بالسرعة عبر إشارات PWM المتمثلة بالطرفين PD4 (OC1B) و PD5 (OC1A) المتصلين بالـ Timer 1 لضبط سرعة السيارة.")

    add_rtl_heading(doc, "مخطط الدائرة وتوصيل العتاد الهيكلي:", level=2)
    add_centered_image(doc, "schmatic5_colorized.png", width_inches=5.5, caption_text="شكل (4): مخطط الدائرة الهندسية التفصيلي لربط ATmega32 مع الحساسات وموديول البلوتوث والمحركات")

    # 11. Conclusion
    add_rtl_heading(doc, "10. الخاتمة والاستنتاجات الفنية (Conclusion)", level=1)
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "نجح مشروع السيارة الذكية رعد في الجمع بين تخصصين هامين في هندسة الحاسبات والتحكم، وهما تطوير التطبيقات البرمجية التفاعلية للهواتف المحمولة وتصميم وبرمجة الأنظمة المدمجة والمتحكمات الدقيقة.")
    
    p = add_rtl_paragraph(doc)
    add_rtl_run(p, "من الناحية البرمجية المتقدمة، تمكنا من تطبيق نماذج متقدمة مثل البرمجة غير المتزامنة لضمان تشغيل واجهة المستخدم وتلقي البيانات من البلوتوث في آن واحد دون تعليق، وتطبيق معالجة الصلاحيات المتقدمة بمستوى أمان عالي، بالإضافة إلى تصميم الواجهات بطابع نيون يعرض البيانات اللحظية والتحذيرات ديناميكياً. ومن الناحية الهندسية، تميز المشروع بوجود نظام حماية مدمج بالعتاد الصلب يضمن حماية السيارة ذاتياً من العقبات المحيطة بها فوراُ.")

    # Save document
    doc_path = "ATmega32_Smart_Car_Advanced_Programming_Report_V3.docx"
    doc.save(doc_path)
    print(f"Success: Report generated at {doc_path}")

if __name__ == "__main__":
    main()
